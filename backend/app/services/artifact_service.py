from __future__ import annotations

from dataclasses import dataclass
from datetime import datetime, timedelta, timezone
from html import escape
import base64
import binascii
import hashlib
from pathlib import Path, PurePosixPath
import json
import re
from typing import Literal

import jwt
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from reportlab.lib.enums import TA_CENTER, TA_LEFT
from reportlab.lib.pagesizes import letter
from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
from reportlab.lib.units import inch
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont
from reportlab.platypus import Paragraph, SimpleDocTemplate, Spacer

from app.core.config import CJK_FONT_NAME, CJK_FONT_PATH, JWT_SECRET, PRIVATE_GENERATED_DIR, PUBLIC_BASE_URL
from app.utils.common import new_id


MAX_DOCUMENT_CHARS = 100_000
MAX_UPLOAD_IMAGE_BYTES = 10 * 1024 * 1024
IMAGE_DATA_URL_RE = re.compile(r"^data:(image/(?:png|jpeg|jpg|webp));base64,(.+)$", re.DOTALL)
ARTIFACT_TTL_DAYS = 30
ARTIFACT_SIGNING_KEY = hashlib.sha256(JWT_SECRET.encode()).digest()
MEDIA_TYPES = {
    ".png": "image/png",
    ".jpg": "image/jpeg",
    ".jpeg": "image/jpeg",
    ".webp": "image/webp",
    ".mp3": "audio/mpeg",
    ".wav": "audio/wav",
    ".m4a": "audio/mp4",
    ".mp4": "video/mp4",
    ".srt": "application/x-subrip",
    ".pdf": "application/pdf",
    ".docx": "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
}
CJK_FONT_CANDIDATES = (
    CJK_FONT_PATH,
    "/System/Library/Fonts/Supplemental/Arial Unicode.ttf",
    "/Library/Fonts/Arial Unicode.ttf",
    "/usr/share/fonts/opentype/noto/NotoSansCJK-Regular.ttc",
    "/usr/share/fonts/truetype/noto/NotoSansCJK-Regular.ttc",
)


@dataclass(frozen=True)
class TextBlock:
    kind: Literal["body", "h1", "h2", "h3", "bullet", "number"]
    text: str
    marker: str = ""


def _plain_inline(value: str) -> str:
    value = re.sub(r"\*\*(.+?)\*\*", r"\1", value)
    value = re.sub(r"__(.+?)__", r"\1", value)
    return re.sub(r"`(.+?)`", r"\1", value).strip()


def parse_markdown_blocks(content: str) -> list[TextBlock]:
    blocks: list[TextBlock] = []
    paragraph: list[str] = []

    def flush() -> None:
        if paragraph:
            blocks.append(TextBlock("body", _plain_inline(" ".join(paragraph))))
            paragraph.clear()

    for raw_line in content.replace("\r\n", "\n").split("\n"):
        line = raw_line.strip()
        if not line:
            flush()
            continue
        heading = re.match(r"^(#{1,3})\s+(.+)$", line)
        bullet = re.match(r"^[-*+]\s+(.+)$", line)
        number = re.match(r"^(\d+)[.)]\s+(.+)$", line)
        if heading:
            flush()
            blocks.append(TextBlock(f"h{len(heading.group(1))}", _plain_inline(heading.group(2))))
        elif bullet:
            flush()
            blocks.append(TextBlock("bullet", _plain_inline(bullet.group(1))))
        elif number:
            flush()
            blocks.append(TextBlock("number", _plain_inline(number.group(2)), number.group(1)))
        else:
            paragraph.append(line)
    flush()
    return blocks


def _validate_document(title: str, content: str) -> tuple[str, str]:
    title, content = title.strip(), content.strip()
    if not title:
        raise ValueError("document title is required")
    if not content:
        raise ValueError("document content is required")
    if len(content) > MAX_DOCUMENT_CHARS:
        raise ValueError(f"document content exceeds {MAX_DOCUMENT_CHARS} characters")
    return title[:160], content


def _set_font(run, size: float, *, bold: bool = False, color: str = "000000") -> None:
    run.font.name = CJK_FONT_NAME
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.color.rgb = RGBColor.from_string(color)
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), CJK_FONT_NAME)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), CJK_FONT_NAME)
    run._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), CJK_FONT_NAME)
    run._element.get_or_add_rPr().rFonts.set(qn("w:cs"), CJK_FONT_NAME)


def _configure_docx(doc: Document) -> None:
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.right_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    tokens = {
        "Normal": (11, "000000", 0, 6),
        "Title": (24, "0B2545", 0, 12),
        "Heading 1": (16, "2E74B5", 16, 8),
        "Heading 2": (13, "2E74B5", 12, 6),
        "Heading 3": (12, "1F4D78", 8, 4),
    }
    for name, (size, color, before, after) in tokens.items():
        style = doc.styles[name]
        style.font.name = CJK_FONT_NAME
        style.font.size = Pt(size)
        style.font.color.rgb = RGBColor.from_string(color)
        for script in ("w:ascii", "w:hAnsi", "w:eastAsia", "w:cs"):
            style._element.get_or_add_rPr().rFonts.set(qn(script), CJK_FONT_NAME)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.line_spacing = 1.1

    for name in ("List Bullet", "List Number"):
        style = doc.styles[name]
        style.font.name = CJK_FONT_NAME
        style.font.size = Pt(11)
        for script in ("w:ascii", "w:hAnsi", "w:eastAsia", "w:cs"):
            style._element.get_or_add_rPr().rFonts.set(qn(script), CJK_FONT_NAME)
        style.paragraph_format.left_indent = Inches(0.5)
        style.paragraph_format.first_line_indent = Inches(-0.25)
        style.paragraph_format.space_after = Pt(8)
        style.paragraph_format.line_spacing = 1.167

    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = footer.add_run("Page ")
    _set_font(run, 9, color="666666")
    field = OxmlElement("w:fldSimple")
    field.set(qn("w:instr"), "PAGE")
    run._element.addnext(field)

    title_ppr = doc.styles["Title"]._element.get_or_add_pPr()
    border = title_ppr.find(qn("w:pBdr"))
    if border is not None:
        title_ppr.remove(border)


def create_docx(path: Path, title: str, content: str) -> None:
    title, content = _validate_document(title, content)
    path.parent.mkdir(parents=True, exist_ok=True)
    doc = Document()
    _configure_docx(doc)
    doc.core_properties.title = title

    paragraph = doc.add_paragraph(style="Title")
    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    title_run = paragraph.add_run(title)
    _set_font(title_run, 24, bold=True, color="0B2545")

    style_for = {"h1": "Heading 1", "h2": "Heading 2", "h3": "Heading 3", "bullet": "List Bullet", "number": "List Number"}
    for block in parse_markdown_blocks(content):
        paragraph = doc.add_paragraph(style=style_for.get(block.kind, "Normal"))
        paragraph.add_run(block.text)
    doc.save(path)


def _pdf_font(text: str) -> str:
    if not any("\u3400" <= char <= "\u9fff" for char in text):
        return "Helvetica"
    path = next((Path(item) for item in CJK_FONT_CANDIDATES if item and Path(item).is_file()), None)
    if path is None:
        raise ValueError("CJK font is unavailable; set SCENEFLOW_CJK_FONT_PATH to a Chinese TTF/TTC font")
    if "SceneFlowCJK" not in pdfmetrics.getRegisteredFontNames():
        pdfmetrics.registerFont(TTFont("SceneFlowCJK", str(path)))
    return "SceneFlowCJK"


def _pdf_styles(font_name: str) -> dict[str, ParagraphStyle]:
    base = getSampleStyleSheet()
    return {
        "title": ParagraphStyle("SceneFlowTitle", parent=base["Title"], fontName=font_name, fontSize=24, leading=31, textColor="#0B2545", alignment=TA_LEFT, spaceAfter=18, wordWrap="CJK"),
        "h1": ParagraphStyle("SceneFlowH1", parent=base["Heading1"], fontName=font_name, fontSize=16, leading=22, textColor="#2E74B5", spaceBefore=16, spaceAfter=8, wordWrap="CJK"),
        "h2": ParagraphStyle("SceneFlowH2", parent=base["Heading2"], fontName=font_name, fontSize=13, leading=19, textColor="#2E74B5", spaceBefore=12, spaceAfter=6, wordWrap="CJK"),
        "h3": ParagraphStyle("SceneFlowH3", parent=base["Heading3"], fontName=font_name, fontSize=12, leading=18, textColor="#1F4D78", spaceBefore=8, spaceAfter=4, wordWrap="CJK"),
        "body": ParagraphStyle("SceneFlowBody", parent=base["BodyText"], fontName=font_name, fontSize=11, leading=15, textColor="#000000", alignment=TA_LEFT, spaceAfter=8, wordWrap="CJK"),
        "list": ParagraphStyle("SceneFlowList", parent=base["BodyText"], fontName=font_name, fontSize=11, leading=15, leftIndent=24, firstLineIndent=-12, spaceAfter=6, wordWrap="CJK"),
        "footer": ParagraphStyle("SceneFlowFooter", parent=base["BodyText"], fontName=font_name, fontSize=9, leading=11, textColor="#666666", alignment=TA_CENTER),
    }


def create_pdf(path: Path, title: str, content: str) -> None:
    title, content = _validate_document(title, content)
    path.parent.mkdir(parents=True, exist_ok=True)
    font_name = _pdf_font(title + content)
    styles = _pdf_styles(font_name)
    story = [Paragraph(escape(title), styles["title"])]
    for block in parse_markdown_blocks(content):
        text = escape(block.text)
        if block.kind == "bullet":
            text = "- " + text
            style = styles["list"]
        elif block.kind == "number":
            text = f"{escape(block.marker)}. {text}"
            style = styles["list"]
        else:
            style = styles.get(block.kind, styles["body"])
        story.append(Paragraph(text, style))
    story.append(Spacer(1, 8))

    def footer(canvas, document) -> None:
        canvas.saveState()
        canvas.setFont(font_name, 9)
        canvas.setFillColor("#666666")
        canvas.drawCentredString(letter[0] / 2, 0.5 * inch, f"Page {document.page}")
        canvas.restoreState()

    SimpleDocTemplate(
        str(path),
        pagesize=letter,
        rightMargin=inch,
        leftMargin=inch,
        topMargin=inch,
        bottomMargin=0.75 * inch,
        title=title,
        author="SceneFlow",
    ).build(story, onFirstPage=footer, onLaterPages=footer)


def _safe_segment(value: str) -> str:
    return re.sub(r"[^A-Za-z0-9_-]", "_", value)[:80] or "chat"


def media_type_for(name: str) -> str:
    return MEDIA_TYPES.get(Path(name).suffix.lower(), "application/octet-stream")


def decode_image_data_url(value: str, maximum: int = MAX_UPLOAD_IMAGE_BYTES) -> tuple[bytes, str, str]:
    """Decode a browser-supplied `data:image/...;base64,` string into bytes, mime, extension.

    The client never uploads multipart — every image the frontend sends (chat attachments,
    image-to-image references, covers, character sheets) arrives as a data URL, so the
    decode and its size ceiling live here once instead of in each caller.

    Raises ValueError, not HTTPException: the API layer decides the status code.
    """
    match = IMAGE_DATA_URL_RE.match(value.strip())
    if not match:
        raise ValueError("image must be a png/jpeg/webp base64 data URL")
    mime_type, encoded = match.groups()
    try:
        data = base64.b64decode(encoded, validate=True)
    except (binascii.Error, ValueError) as exc:
        raise ValueError("invalid base64 image data") from exc
    if not data or len(data) > maximum:
        raise ValueError(f"image must be 1 byte to {maximum // (1024 * 1024)}MB")
    extension = "jpg" if mime_type in {"image/jpeg", "image/jpg"} else mime_type.removeprefix("image/")
    return data, mime_type, extension


def _artifact_path(scope: str, extension: str, category: str = "chat") -> Path:
    # ponytail: local disk is enough for one backend instance; move to object storage when deployments become multi-instance.
    directory = PRIVATE_GENERATED_DIR / _safe_segment(category) / _safe_segment(scope)
    directory.mkdir(parents=True, exist_ok=True)
    return directory / f"{new_id('artifact')}.{extension}"


def artifact_relative_path(path: Path) -> str:
    """Express an on-disk artifact as a path relative to the private artifact root."""
    return path.resolve().relative_to(PRIVATE_GENERATED_DIR.resolve()).as_posix()


def artifact_absolute_path(relative: str) -> Path:
    """Resolve a stored relative path back to disk, rejecting escapes from the artifact root."""
    pure = PurePosixPath(relative)
    if not relative or pure.is_absolute() or ".." in pure.parts:
        raise ValueError("invalid artifact path")
    path = (PRIVATE_GENERATED_DIR / Path(*pure.parts)).resolve()
    path.relative_to(PRIVATE_GENERATED_DIR.resolve())
    return path


def _signed_url(path: Path, filename: str, media_type: str, inline: bool) -> str:
    return _sign(artifact_relative_path(path), filename, media_type, inline)


def _sign(relative: str, filename: str, media_type: str, inline: bool) -> str:
    issued = datetime.now(timezone.utc)
    token = jwt.encode(
        {
            "scope": "artifact",
            "path": relative,
            "filename": filename,
            "mediaType": media_type,
            "inline": inline,
            "iat": issued,
            "exp": issued + timedelta(days=ARTIFACT_TTL_DAYS),
        },
        ARTIFACT_SIGNING_KEY,
        algorithm="HS256",
    )
    return f"{PUBLIC_BASE_URL}/api/chat/artifacts/{token}"


def signed_file_url(path: Path, filename: str, media_type: str, inline: bool = True) -> str:
    return _signed_url(path, Path(filename).name[:180], media_type[:120], inline)


def signed_url_for_stored(relative: str, download_stem: str = "", inline: bool = True) -> str:
    """Mint a fresh link for an artifact the database tracks by path.

    Signed links expire after ARTIFACT_TTL_DAYS. Persisting the URL would make every
    asset in a long-running series 404 a month later, so rows keep the relative path and
    the URL is minted per response instead. The extension comes from the stored file, so a
    caller naming the download cannot mislabel a webp as a png.
    """
    artifact_absolute_path(relative)
    stored = PurePosixPath(relative)
    stem = Path(download_stem).name[:180]
    name = f"{stem}{stored.suffix}" if stem else stored.name
    return _sign(relative, name, media_type_for(relative)[:120], inline)


def stored_relative_path(value: str) -> str:
    """Recover the relative path a legacy signed URL points at, ignoring its expiry.

    Used once by the schema migration that moves rows off stored URLs; an unreadable
    token means the link was already dead and the caller drops the reference.
    """
    token = value.rsplit("/", 1)[-1].strip()
    payload = jwt.decode(token, ARTIFACT_SIGNING_KEY, algorithms=["HS256"], options={"verify_exp": False})
    if payload.get("scope") != "artifact":
        raise ValueError("invalid scope")
    relative = str(payload["path"])
    artifact_absolute_path(relative)
    return relative


def save_binary_artifact(scope: str, filename: str, data: bytes, media_type: str, inline: bool = True) -> str:
    extension = Path(filename).suffix.lower().removeprefix(".") or "bin"
    path = _artifact_path(scope, extension)
    path.write_bytes(data)
    path.chmod(0o600)
    return signed_file_url(path, filename, media_type, inline)


def store_artifact(category: str, scope: str, filename: str, data: bytes) -> str:
    """Persist bytes under the private root and return the relative path to store in a row."""
    path = _artifact_path(scope, Path(filename).suffix.lower().removeprefix(".") or "bin", category)
    path.write_bytes(data)
    path.chmod(0o600)
    return artifact_relative_path(path)


def artifact_from_token(token: str) -> tuple[Path, str, str, bool]:
    try:
        payload = jwt.decode(token, ARTIFACT_SIGNING_KEY, algorithms=["HS256"])
        if payload.get("scope") != "artifact":
            raise ValueError("invalid scope")
        path = artifact_absolute_path(str(payload["path"]))
        if not path.is_file():
            raise FileNotFoundError(path)
        return path, str(payload["filename"])[:180], str(payload["mediaType"])[:120], bool(payload.get("inline"))
    except Exception as exc:
        raise ValueError("invalid or expired artifact link") from exc


def _download_name(title: str, extension: str) -> str:
    stem = re.sub(r"[\\/:*?\"<>|\r\n]+", " ", title).strip(" .")[:80] or "SceneFlow"
    return f"{stem}.{extension}"


def _markdown_label(value: str) -> str:
    return re.sub(r"[\r\n]+", " ", value).replace("\\", "\\\\").replace("[", "\\[").replace("]", "\\]")


def save_image_artifact(session_id: str, title: str, data: bytes, extension: str) -> dict[str, str]:
    extension = "jpg" if extension.lower() in {"jpeg", "jpg"} else extension.lower()
    if extension not in {"png", "jpg", "webp"}:
        extension = "png"
    path = _artifact_path(session_id, extension)
    path.write_bytes(data)
    path.chmod(0o600)
    media_type = "image/jpeg" if extension == "jpg" else f"image/{extension}"
    filename = _download_name(title or "generated-image", extension)
    url = _signed_url(path, filename, media_type, True)
    return {"kind": "image", "url": url, "filename": filename, "markdown": f"![{_markdown_label(title or '生成图片')}]({url})"}


def save_document_artifact(session_id: str, kind: Literal["pdf", "docx"], title: str, content: str) -> dict[str, str]:
    path = _artifact_path(session_id, kind)
    if kind == "pdf":
        create_pdf(path, title, content)
        media_type, label, inline = "application/pdf", "下载 PDF", True
    else:
        create_docx(path, title, content)
        media_type, label, inline = "application/vnd.openxmlformats-officedocument.wordprocessingml.document", "下载 Word 文档", False
    path.chmod(0o600)
    filename = _download_name(title, kind)
    url = _signed_url(path, filename, media_type, inline)
    return {"kind": kind, "url": url, "filename": filename, "markdown": f"[{_markdown_label(f'{label}：{title}')}]({url})"}


def tool_result(payload: dict[str, str]) -> str:
    return json.dumps(payload, ensure_ascii=False)
